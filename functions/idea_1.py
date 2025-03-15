from openai import AsyncOpenAI, AsyncAzureOpenAI
import openai
import json
import os
import asyncio

# import os
from docx import Document
from markdownify import markdownify as md
import re

from dotenv import load_dotenv
load_dotenv("/agent_vdo_mini/.env") ## load from intersys directory
load_dotenv("./.env") ## load from current directory
load_dotenv("../.env") ## load from parent directory


api_key = os.environ.get("OPENAI_API_KEY_PTU")
endpoint = os.environ.get("OPENAI_API_BASE_PTU")
deployment_name = os.environ.get("OPENAI_DEPLOYMENT_NAME_PTU")
api_version = os.environ.get("OPENAI_API_VERSION_PTU")
if all([api_key, endpoint, deployment_name, api_version]):
    client = AsyncAzureOpenAI(api_key=api_key, azure_endpoint=endpoint, api_version=api_version)
else:
    print("Missing environment variables for Azure, use default OpenAI Instead")
    deployment_name = os.environ.get("OPENAI_DEPLOYMENT_NAME")
    api_key = os.environ.get("OPENAI_API_KEY")
    endpoint = os.environ.get("OPENAI_API_BASE")
    try:
        client = AsyncOpenAI(api_key=api_key, base_url=endpoint)
    except Exception as e:
        print(e)
    # if all([deployment_name, api_key, endpoint]):
    #     client = AsyncOpenAI(api_key=api_key, base_url=endpoint)
    # else: 
    #     raise ValueError("Missing environment variables for OpenAI")


async def a_get_bot_response(user_input, system_prompt, chat_history=None, retry=0, number_of_conversations=0):
    if chat_history is None:
        chat_history = []
    try:
        results = await client.chat.completions.create(
            model=deployment_name,
            messages = [
                {
                    "role": "system", 
                    "content": system_prompt
                },
                *chat_history[number_of_conversations*-2:],
                {"role": "user", "content": user_input}
            ]
        )
        return results.choices[0].message.content
    except openai.APIConnectionError:
        if retry < 5:
            print(f"Connection error. Retrying {retry+1}")
            return await a_get_bot_response(user_input, system_prompt, chat_history, retry=retry+1, number_of_conversations=number_of_conversations)
        return f"Timeout Error. Retried {retry} times."


def extract_json(string):
    string = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', string)
    string = re.sub(r'(?<!\\)"((?:\\.|[^"\\])*)"', lambda m: '"{}"'.format(m.group(1).replace('\n', '\\n')), string) ## prevents multi-line responses
    start_index = string.find('{')
    end_index = string.rfind('}')
    string = string[start_index:end_index + 1]
    string = json.loads(string)
    return string

def extract_backticks(string, prefixes=["python", "json"]):
    start_index = string.find('```')
    end_index = string.rfind('```')
    if start_index == -1 or end_index == -1:
        print("No backticks found")
        return string
    elif start_index == end_index:
        raise ValueError("Only one backtick found")
    string = string[start_index+3:end_index]
    for prefix in prefixes:
        if string.startswith(prefix):
            string = string[len(prefix):]
            break
    return string            
    

def extract_eval_json(string):
    string = extract_backticks(string)
    if type(string)!=dict:
        try:
            final = extract_json(string)
        except:
            final = eval(string)
    else:
        final = string
    return final

## FOR NIC 
def extract_keywords_from_inputs(inputs):
    extract_keywords_system_promt = """\
You are an expert medical practitioner that extracts key patient-reported outcomes from clinical notes. Your task is to first identify and retrieve terms related to the EQ5D metrics from the provided clinical notes. The EQ5D metrics include the following categories: MOBILITY, SELF-CARE, USUAL ACTIVITIES, PAIN / DISCOMFORT, and ANXIETY / DEPRESSION. Each category is scored from 1 to 5 based on the patient's reported outcomes.

EQ5D Metrics:
- MOBILITY
  1 - I have no problems in walking about
  2 - I have slight problems in walking about
  3 - I have moderate problems in walking about
  4 - I have severe problems in walking about
  5 - I am unable to walk about

- SELF-CARE
  1 - I have no problems washing or dressing myself
  2 - I have slight problems washing or dressing myself
  3 - I have moderate problems washing or dressing myself
  4 - I have severe problems washing or dressing myself
  5 - I am unable to wash or dress myself

- USUAL ACTIVITIES (e.g. work, study, housework, family, or leisure activities)
  1 - I have no problems doing my usual activities
  2 - I have slight problems doing my usual activities
  3 - I have moderate problems doing my usual activities
  4 - I have severe problems doing my usual activities
  5 - I am unable to do my usual activities

- PAIN / DISCOMFORT
  1 - I have no pain or discomfort
  2 - I have slight pain or discomfort
  3 - I have moderate pain or discomfort
  4 - I have severe pain or discomfort
  5 - I have extreme pain or discomfort

- ANXIETY / DEPRESSION
  1 - I am not anxious or depressed
  2 - I am slightly anxious or depressed
  3 - I am moderately anxious or depressed
  4 - I am severely anxious or depressed
  5 - I am extremely anxious or depressed

Input: The input will be the doctor's notes, which can be admission, discharge, or daily encounter notes. The information extracted for EQ5D must be based on questions asked by the doctor to the patient on the same day, not derived from previous complaints.

Output: The output should be in JSON format, with keys for each EQ5D category. Each key should have a value as a comma-separated string containing all the related terms found in the document.

Example JSON format:
{
  "mobility": "walking, walking about, problems in walking",
  "self_care": "washing, dressing, problems washing, problems dressing",
  "usual_activities": "work, study, housework, family, leisure activities, problems doing usual activities",
  "pain_discomfort": "pain, discomfort, slight pain, severe pain",
  "anxiety_depression": "anxious, depressed, slight anxiety, severe depression"
}

Note: If no related terms are found for a category, the value should be an empty string.
"""
    user_input = inputs
    chat_history = []
    number_of_conversations = 0
    response = asyncio.run(
        a_get_bot_response(
            user_input, 
            system_prompt=extract_keywords_system_promt, 
            chat_history=chat_history, 
            number_of_conversations=number_of_conversations,
        )
    )
    try:
        # json.loads(response)
        response = extract_eval_json(response)
    except json.JSONDecodeError:
        print(response)
        raise ValueError("Response is not in JSON format")
    return response

from functions.faiss_generator import getIndex, compareQuery
def runRetrieval(keywords, top_k=2):
    if type(keywords) == str:
        keywords = [keywords,]
    keywordstr = " ".join(keywords)
    index = getIndex()
    # contexts = compareQuery(index, keywordstr, k=top_k)
    texts, labels, scores = compareQuery(index, keywordstr, k=top_k)
    return texts, labels, scores

# ---------------------------------------------

def extract_eq5d(markdown_content):
    
    start_index = markdown_content.lower().find("eq-5d scores".lower())

    if start_index == -1:
        print("No eq5d found in the document")
        return ""
    
    end_index = markdown_content.lower().find("issues and progress".lower())

    if end_index == -1:
        end_index = markdown_content.lower().find("issues & progress".lower())
        if end_index == -1:
            end_index = markdown_content.lower().find("**Issues & Progress**".lower())
            if end_index == -1:
                print("No issues and progress found in the document")
                return ""
    markdown_content = markdown_content[start_index:end_index+1]
        
    return markdown_content

def extract_issues(markdown_content):
    
    start_index = markdown_content.lower().find("issues and progress".lower())

    if start_index == -1:
        start_index = markdown_content.lower().find("issues & progress".lower())
        if start_index == -1:
            print("No issues and progress found in the document")
            return ""
    markdown_content = markdown_content[start_index:]
        
    return markdown_content

def extract_soap(markdown_content):
    
    end_index = markdown_content.lower().find("**Objective".lower())

    if end_index == -1:
        print("No '**Objective' found in the document")
    markdown_content = markdown_content[:end_index]
    
    current_admission_index = markdown_content.lower().find("**Current Admission".lower())
    subjective_index = markdown_content.lower().find("**Subjective".lower())
    if current_admission_index != -1:
        markdown_content = markdown_content[current_admission_index:]
    elif subjective_index != -1:
        markdown_content = markdown_content[subjective_index:]
    else:
        print("No '**Subjective' or '**Current Admission' found in the document")
        
    return markdown_content


# # Load the Word document
def get_text_from_docx_raw(doc_path):
    try:
        doc = Document(doc_path)
    except:
        doc = Document("/agent_vdo_mini/"+doc_path)
        
    # Initialize an empty string to hold the Markdown content
    markdown_content = ""

    def convert_run_to_markdown(run):
        text = run.text
        if run.bold:
            text = f"**{text}**"
        if run.italic:
            text = f"*{text}*"
        return text

    def get_paragraph_numbering(para):
        numbering = para._element.xpath('w:pPr/w:numPr')
        if numbering:
            num_id = numbering[0].xpath('w:numId/@w:val')[0]
            ilvl = numbering[0].xpath('w:ilvl/@w:val')[0]
            return int(num_id), int(ilvl)
        return None, None

    # Iterate through the paragraphs in the document and convert them to Markdown
    for para in doc.paragraphs:
        para_text = ""
        for run in para.runs:
            para_text += convert_run_to_markdown(run)

        # Check for headings
        if para.style.name.startswith('Heading'):
            level = int(para.style.name.split()[-1])
            para_text = f"{'#' * level} {para_text}"
        # Check for numbered list items
        num_id, ilvl = get_paragraph_numbering(para)
        if num_id is not None:
            # This is a numbered list item
            para_text = f"{ilvl + 1}. {para_text.strip()}"
        elif para.style.name.startswith('List Number'):
            print("LIST NUMBER")
            # Extract the list number from the paragraph text
            list_number = para.text.split('.')[0]
            para_text = f"{list_number}. {para_text[len(list_number)+1:].strip()}"
        # Check for bullet list items
        elif para.style.name.startswith('List Bullet'):
            para_text = f"- {para_text.strip()}"
        # Check for block quotes
        # elif para.style.name == 'Quote':
        #     para_text = f"> {para_text.strip()}"

        markdown_content += para_text + "\n\n"

    return markdown_content

# # Load the Word document
def get_text_from_docx(doc_path):
    try:
        doc = Document(doc_path)
    except:
        doc = Document("/agent_vdo_mini/"+doc_path)
        
    # Initialize an empty string to hold the Markdown content
    markdown_content = ""

    def convert_run_to_markdown(run):
        text = run.text
        if run.bold:
            text = f"**{text}**"
        if run.italic:
            text = f"*{text}*"
        return text

    def get_paragraph_numbering(para):
        numbering = para._element.xpath('w:pPr/w:numPr')
        if numbering:
            num_id = numbering[0].xpath('w:numId/@w:val')[0]
            ilvl = numbering[0].xpath('w:ilvl/@w:val')[0]
            return int(num_id), int(ilvl)
        return None, None

    # Iterate through the paragraphs in the document and convert them to Markdown
    for para in doc.paragraphs:
        para_text = ""
        for run in para.runs:
            para_text += convert_run_to_markdown(run)

        # Check for headings
        if para.style.name.startswith('Heading'):
            level = int(para.style.name.split()[-1])
            para_text = f"{'#' * level} {para_text}"
        # Check for numbered list items
        num_id, ilvl = get_paragraph_numbering(para)
        if num_id is not None:
            # This is a numbered list item
            para_text = f"{ilvl + 1}. {para_text.strip()}"
        elif para.style.name.startswith('List Number'):
            print("LIST NUMBER")
            # Extract the list number from the paragraph text
            list_number = para.text.split('.')[0]
            para_text = f"{list_number}. {para_text[len(list_number)+1:].strip()}"
        # Check for bullet list items
        elif para.style.name.startswith('List Bullet'):
            para_text = f"- {para_text.strip()}"
        # Check for block quotes
        # elif para.style.name == 'Quote':
        #     para_text = f"> {para_text.strip()}"

        markdown_content += para_text + "\n\n"

    markdown_content = extract_soap(markdown_content)
    return markdown_content

# # Load the Word document
def get_text_from_discharge(doc_path):
    markdown_content = extract_issues(get_text_from_docx_raw(doc_path))
    return markdown_content

def extract_prom_prem_from_text(inputs, use_rag=False):
    system_prompt = """
You are an expert medical practitioner that extracts key patient-reported outcomes from clinical notes. Your task is to analyze the provided clinical notes and extract information related to the EQ5D metrics. The EQ5D metrics include the following categories: MOBILITY, SELF-CARE, USUAL ACTIVITIES, PAIN / DISCOMFORT, and ANXIETY / DEPRESSION. Each category is scored from 1 to 5 based on the patient's reported outcomes.

EQ5D Metrics:
- MOBILITY
  1 - I have no problems in walking about
  2 - I have slight problems in walking about
  3 - I have moderate problems in walking about
  4 - I have severe problems in walking about
  5 - I am unable to walk about

- SELF-CARE
  1 - I have no problems washing or dressing myself
  2 - I have slight problems washing or dressing myself
  3 - I have moderate problems washing or dressing myself
  4 - I have severe problems washing or dressing myself
  5 - I am unable to wash or dress myself

- USUAL ACTIVITIES (e.g. work, study, housework, family, or leisure activities)
  1 - I have no problems doing my usual activities
  2 - I have slight problems doing my usual activities
  3 - I have moderate problems doing my usual activities
  4 - I have severe problems doing my usual activities
  5 - I am unable to do my usual activities

- PAIN / DISCOMFORT
  1 - I have no pain or discomfort
  2 - I have slight pain or discomfort
  3 - I have moderate pain or discomfort
  4 - I have severe pain or discomfort
  5 - I have extreme pain or discomfort

- ANXIETY / DEPRESSION
  1 - I am not anxious or depressed
  2 - I am slightly anxious or depressed
  3 - I am moderately anxious or depressed
  4 - I am severely anxious or depressed
  5 - I am extremely anxious or depressed

Input: The input will be the doctor's notes, which can be admission, discharge, or daily encounter notes. The information extracted for EQ5D must be based on questions asked by the doctor to the patient on the same day, not derived from previous complaints.

Output: The output should be in JSON format, with keys for each category in the main JSON. Each category should have a sub-JSON with the following keys:
- reason_for_score: a string explaining the reason for the score
- supporting_statements: a list of strings containing supporting statements from the document
- captured_in_doc: a boolean indicating if the information was captured in the document (true/false)
- final_score: an integer from 1 to 5, or null if there is no mention of tracking by the doctor

Example JSON format for ANXIETY / DEPRESSION:
{
  "anxiety_depression": {
    "reason_for_score": "Patient reports slight anxiety but no severe depression.",
    "supporting_statements": [
      "reports feeling slightly anxious",
      "No signs of severe depression observed."
    ],
    "captured_in_doc": true,
    "final_score": 2
  },
  "mobility": { ... },
  "self_care": { ... },
  "usual_activities": { ... },
  "pain_discomfort": { ... }
}

For the value captured in `supporting_statements`, they must be EXACT QUOTATIONS from the document. If there is a typographical mistake or incorrect information, include it as it appears in the original text without any corrections, to indicate that it is an exact reproduction.

------------------

Note: If the doctor did not capture information based on the rubrics given, ensure the LLM follows the rubrics closely. For example, if the doctor asked about balance but did not ask about mobility/walking about, mark it as null with the reason for score as "not captured."

Note: Do not use the "Premorbid ADL" in any of the assessments, as it was not something asked by the doctor. Only use patient reported outcomes based on the EQ5D metrics.

Note: For anxiety and depression, the doctor can also ask about the mood. If the report contains information about the patient's mood, include it in the supporting statements.

**IMPORTANT**
You MUST NOT guess or infer. If the patient did not report the information, do not include it in the output. The information must be directly from the patient's response to the doctor.

"""
    ## additional prompts that failed
    """\
# More information for each EQ-5D
For Usual Activities, the doctor can ask about the patient's ability to perform work, study, housework, family, or leisure activities.
Responses considered under usual activities in EQ5D include work-related tasks, household chores, family and social activities, hobbies, personal care, and educational activities.
Responses NOT considered under usual activities in EQ5D include activities that are not part of the individual's regular routine, unusual or one-time events, tasks not typically performed by the individual, irrelevant tasks, and activities beyond the individual's capacity. Financial stability-related answers such as payment of rents, bills, or loans are also not considered under usual activities.
"""
    user_input = inputs
    chat_history = []
    if use_rag:
        rag_contents = []
        keywords = extract_keywords_from_inputs(user_input)
        for k, v in keywords.items():
            if v:
                retrieved_results = runRetrieval(v, top_k=2)
                texts, labels, scores = retrieved_results
                for i, (text, label, score) in enumerate(zip(texts, labels, scores)):
                    rag_contents.append(f"### {label.upper()}\n- {text}\n - Score: {score}")
            
        rag_content = "\n\n".join(rag_contents)
        
        chat_history = [
            {
                "role": "user",
                "content": "Here is some information that will be useful for you:\n\n{rag_content}".format(rag_content=rag_content)
            }
        ]
    number_of_conversations = 0
    ## TODO: change to async
    response = asyncio.run(
        a_get_bot_response(
            user_input, 
            system_prompt, 
            chat_history, 
            number_of_conversations=number_of_conversations,
            )
    )
    
    return response

def extract_prom_prem_from_docx(doc_path, use_rag=False):
    text = get_text_from_docx(doc_path)
    text = extract_soap(text)
    print(text)
    return extract_prom_prem_from_text(text, use_rag=use_rag)

def extract_eortc_from_text(inputs):
    system_prompt = """\
You are an expert medical practitioner that extracts key patient-reported outcomes from clinical notes. Your task is to analyze the provided clinical notes and extract information related to the EORTC metrics. The EORTC metrics include the following categories: PHYSICAL FUNCTIONING, ROLE FUNCTIONING, EMOTIONAL FUNCTIONING, COGNITIVE FUNCTIONING, SOCIAL FUNCTIONING, and GLOBAL HEALTH STATUS/QUALITY OF LIFE. Each category is scored based on the patient's reported outcomes.

EORTC Metrics:
1. Do you have any trouble doing strenuous activities, like carrying a heavy shopping bag or a suitcase?
    - Not at all - 1
    - A little - 2
    - Quite a bit - 3
    - Very much - 4
2. Do you have any trouble taking a long walk?
    - Not at all - 1
    - A little - 2
    - Quite a bit - 3
    - Very much - 4
3. Do you have any trouble taking a short walk outside of the house?
    - Not at all - 1
    - A little - 2
    - Quite a bit - 3
    - Very much - 4
4. Do you need to stay in bed or a chair during the day?
    - Not at all - 1
    - A little - 2
    - Quite a bit - 3
    - Very much - 4
5. Do you need help with eating, dressing, washing yourself, or using the toilet?
    - Not at all - 1
    - A little - 2
    - Quite a bit - 3
    - Very much - 4
6. Were you limited in doing either your work or other daily activities?
    - Not at all - 1
    - A little - 2
    - Quite a bit - 3
    - Very much - 4
7. Were you limited in pursuing your hobbies or other leisure time activities?
    - Not at all - 1
    - A little - 2
    - Quite a bit - 3
    - Very much - 4
8. Have you had difficulty in concentrating on things, like reading a newspaper or watching television?
    - Not at all - 1
    - A little - 2
    - Quite a bit - 3
    - Very much - 4
9. Did you feel tense?
    - Not at all - 1
    - A little - 2
    - Quite a bit - 3
    - Very much - 4
10. Did you worry?
    - Not at all - 1
    - A little - 2
    - Quite a bit - 3
    - Very much - 4
11. Did you feel irritable?
    - Not at all - 1
    - A little - 2
    - Quite a bit - 3
    - Very much - 4
12. Did you feel depressed?
    - Not at all - 1
    - A little - 2
    - Quite a bit - 3
    - Very much - 4
13. Have you had difficulty remembering things?
    - Not at all - 1
    - A little - 2
    - Quite a bit - 3
    - Very much - 4
14. Has your physical condition or medical treatment interfered with your family life?
    - Not at all - 1
    - A little - 2
    - Quite a bit - 3
    - Very much - 4
15. Has your physical condition or medical treatment interfered with your social activities?
    - Not at all - 1
    - A little - 2
    - Quite a bit - 3
    - Very much - 4
16. How would you rate your overall health during the past week?
    - Very poor - 1
    - 2
    - 3
    - 4
    - 5
    - 6
    - Excellent - 7
17. How would you rate your overall quality of life during the past week?
    - Very poor - 1
    - 2
    - 3
    - 4
    - 5
    - 6
    - Excellent - 7

For the last two questions "16. How would you rate your overall health during the past week?" and "17. How would you rate your overall quality of life during the past week?", **DO NOT INFER THE RESPONSE QUESTIONS. ONLY PICK UP IF IT'S REPORTED EXACTLY AS SUCH IN THE CASE NOTES**

Input: The input will be the doctor's notes, which can be admission, discharge, or daily encounter notes. The information extracted for EORTC must be based on questions asked by the doctor to the patient on the same day, not derived from previous complaints.

Output: The output should be in JSON format, with keys for each question in the main JSON. Each question should have a sub-JSON with the following keys:
- reason_for_score: a string explaining the reason for the score
- supporting_statements: a list of strings containing supporting statements from the document
- captured_in_doc: a boolean indicating if the information was captured in the document (true/false)
- final_score: an integer from 1 to 4 (or 1 to 7 for GLOBAL HEALTH STATUS/QUALITY OF LIFE), or null if there is no mention of tracking by the doctor

Example JSON format output:
{
  "q1_trouble_doing_strenuous_activities": {
        "reason_for_score": "Patient reports having trouble carrying heavy items.",
        "supporting_statements": [
        "Patient reports having trouble carrying heavy items."
        ],
        "captured_in_doc": true,
        "final_score": 3
    },
  "q2_trouble_taking_long_walk": { ... },
  "q3_trouble_taking_short_walk": { ... },
  "q4_need_to_stay_in_bed_or_chair": { ... },
  "q5_need_help_with_eating_dressing_washing_toilet": { ... },
  "q6_limited_in_work_or_daily_activities": { ... },
  "q7_limited_in_hobbies_or_leisure_activities": { ... },
  "q8_difficulty_concentrating": { ... },
  "q9_feel_tense": { ... },
  "q10_worry": { ... },
  "q11_feel_irritable": { ... },
  "q12_feel_depressed": { ... },
  "q13_difficulty_remembering": { ... },
  "q14_interference_with_family_life": { ... },
  "q15_interference_with_social_activities": { ... },
  "q16_overall_health_rating": { ... },
  "q17_overall_quality_of_life_rating": { ... }
}

For the value captured in `supporting_statements`, they must be EXACT QUOTATIONS from the document. If there is a typographical mistake or incorrect information, include it as it appears in the original text without any corrections, to indicate that it is an exact reproduction.

------------------

Note: If the doctor did not capture information based on the rubrics given, ensure the LLM follows the rubrics closely. For example, if the doctor asked about balance but did not ask about physical functioning, mark it as null with the reason for score as "not captured."

Note: Do not use the "Premorbid ADL" in any of the assessments, as it was not something asked by the doctor. Only use patient reported outcomes based on the EORTC metrics.

**IMPORTANT**
You MUST NOT guess or infer. If the patient did not report the information, do not include it in the output. The information must be directly from the patient's response to the doctor.
"""
    user_input = inputs
    chat_history = []
    number_of_conversations = 0
    response = asyncio.run(
        a_get_bot_response(
            user_input, 
            system_prompt, 
            chat_history, 
            number_of_conversations=number_of_conversations,
            )
    )
    
    return response

def extract_eortc_from_docx(doc_path):
    text = get_text_from_docx(doc_path)
    text = extract_soap(text)
    return extract_eortc_from_text(text)



if __name__ == "__main__":
    print("Running the script")    
    # result = extract_prom_prem_from_docx("eq5d/Case 2 admission note (perfect).docx")
    result = extract_prom_prem_from_docx("eq5d/Live case 1 deidentified-synthetic.docx")
    print(result)
    